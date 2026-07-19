from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT = PROJECT_ROOT / "build" / "AI产品项目标准推进SOP.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

FONT = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(96, 96, 96)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
BORDER = "D9E2EC"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_paragraph(doc, text="", size=11, color=None, bold=False, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para_spacing(p, before=before, after=after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    set_para_spacing(p, before=before, after=after)
    p.style = f"Heading {min(level, 3)}"
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_rule(paragraph, color="D7DBE2", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=9.5, color=None, fill=None, align=None):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    set_para_spacing(p, after=0, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_twips = sum(int(Cm(w).twips) for w in widths_cm)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width).twips)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(Cm(width).twips)))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_cm, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    set_table_width(table, widths_cm)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    set_row_cant_split(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, size=9.5, color=NAVY, fill=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        table_row = table.add_row()
        set_row_cant_split(table_row)
        cells = table_row.cells
        for i, value in enumerate(row):
            fill = WHITE if len(table.rows) % 2 else "FBFCFE"
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, fill=fill, align=align)
    add_paragraph(doc, "", after=2)
    return table


def add_callout(doc, title, body, trailing_space=True):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="CAD6E2", size="6")
    set_table_width(table, [16.5])
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_BLUE)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=3, line=1.15)
    r = p.add_run(title)
    set_run_font(r, size=11, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    set_para_spacing(p2, after=0, line=1.15)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    if trailing_space:
        add_paragraph(doc, "", after=2)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("AI 产品项目 SOP")
    set_run_font(run, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("AI 产品项目协作参考")
    set_run_font(r, size=9, color=GRAY)

    return doc


def build_doc():
    doc = setup_document()

    add_paragraph(doc, "项目协作指南", size=11, color=BLUE, bold=True, after=6)
    title = add_paragraph(doc, "AI 产品项目标准推进 SOP", size=24, color=NAVY, bold=True, after=4)
    subtitle = add_paragraph(
        doc,
        "适用于从需求立项、方案设计、Agent 开发、测试验证到上线复盘的完整 AI 产品推进流程",
        size=12.5,
        color=GRAY,
        after=14,
    )
    add_rule(subtitle)

    meta_rows = [
        ("适用对象", "产品经理、前端开发、后端开发、Agent/算法开发、测试、运营及项目协作方"),
        ("适用项目", "AI 文档生成、AI 工作流与智能体应用等 AI 能力建设"),
        ("使用方式", "作为项目启动、研发评审、测试验收、上线复盘的统一流程框架"),
    ]
    add_table(doc, ["项目", "说明"], meta_rows, [3.0, 13.5], font_size=10)

    add_callout(
        doc,
        "核心共识",
        "AI 产品项目不是普通功能开发外接一个模型接口，而是同时管理产品体验、数据上下文、Agent 行为、工程稳定性和效果评测的协作过程。只有这五件事都形成闭环，才算真正完成一个 AI 产品项目。",
    )

    add_heading(doc, "一、总流程", 1)
    add_paragraph(
        doc,
        "建议将 AI 产品项目统一拆成 7 个阶段。前 4 个阶段解决“做什么、为什么做、怎么做”的问题，后 3 个阶段解决“如何实现、如何验证、如何持续优化”的问题。",
        after=8,
    )
    overview_rows = [
        ("1", "需求分析与价值初判", "判断需求是否真实、是否值得做、是否适合用 AI 做", "需求分析文档、价值判断、AI 适配度结论", "场景真实、价值明确、AI 有必要性"),
        ("2", "市场与竞品调研", "明确外部解法、体验基准和差异化方向", "竞品调研报告、可借鉴点、差异化机会", "知道别人怎么做，也知道我们为什么这样做"),
        ("3", "产品方案设计与数据准备", "把需求转化为可开发、可验证的产品方案", "原型、PRD、AI_Spec、数据清单、埋点方案", "产品流程清楚，数据边界清楚，开发能理解并评估"),
        ("4", "Agent 设计与技术方案评审", "确认 AI 能力如何实现，以及工程上是否可行", "Agent 设计文档、技术方案、接口文档、排期评估", "方案可实现，风险可控，成本可接受"),
        ("5", "产品前后端开发 + Agent 开发", "完成功能闭环和 AI 能力实现", "可测试版本、接口实现、Agent 配置、Prompt 版本记录", "主流程跑通，AI 能力可调用，关键日志可追踪"),
        ("6", "功能测试及数据验证", "验证功能可用、AI 输出可信、数据安全", "测试报告、AI 评测表、问题清单、上线建议", "功能稳定，AI 输出达标，无高风险问题"),
        ("7", "上线、监控与复盘", "观察真实使用效果，并沉淀经验", "上线记录、效果看板、用户反馈、复盘报告", "有数据可看，有问题可追，有经验可复用"),
    ]
    add_table(doc, ["阶段", "名称", "核心目标", "主要产出", "通过标准"], overview_rows, [1.1, 3.3, 4.1, 4.0, 4.0], font_size=8.3)

    add_heading(doc, "二、阶段说明", 1)

    stages = [
        (
            "1. 需求分析与价值初判",
            "判断该需求是否值得投入，以及 AI 是否是更合适的解决方式。",
            "明确目标用户、核心场景、痛点和触发频率；判断传统规则、模板、人工流程是否已经足够；评估 AI 能带来的效率、质量或规模化价值；识别失败成本和人工确认要求。",
            "需求分析文档、价值判断、AI 适配度结论、优先级建议。",
            "需求来自真实业务场景；收益与投入基本匹配；AI 相比非 AI 方案有明确增量价值。",
        ),
        (
            "2. 市场与竞品调研",
            "明确行业已有解法，建立体验基准，并找到适合目标业务的差异化方向。",
            "调研同类产品和行业 AI 应用；拆解竞品入口、交互方式、能力边界、结果呈现、收费方式和用户反馈；总结可借鉴点、不可照搬点和差异化机会。",
            "竞品调研报告、体验截图或流程拆解、可借鉴点、差异化机会。",
            "团队知道外部方案怎么做，也能解释我们为什么采用当前方案。",
        ),
        (
            "3. 产品方案设计与数据准备",
            "将需求转化为开发可理解、测试可验证、上线可监控的产品方案。",
            "设计用户流程、页面原型和 PRD；明确 AI 输入输出、数据来源、权限范围、埋点和测试数据；设计用户确认、编辑、撤销、重试和失败兜底机制。",
            "原型、PRD、AI_Spec、数据清单、权限说明、埋点方案、测试样例初稿。",
            "产品流程闭环清楚；数据来源和边界清楚；AI 输出形态和验收口径清楚。",
        ),
        (
            "4. Agent 设计与技术方案评审",
            "确认 Agent 如何工作，以及前后端、数据、权限、成本和稳定性是否可行。",
            "设计 Agent 工作流、Prompt、模型选择、工具调用、输出 schema 和异常处理；评审接口、数据处理、权限校验、日志、灰度、开关、回滚和成本估算。",
            "Agent 设计文档、技术方案、接口文档、模型与工具调用方案、排期评估。",
            "实现路径明确；关键风险有处理方案；成本、性能和权限边界可接受。",
        ),
        (
            "5. 产品前后端开发 + Agent 开发",
            "完成产品功能与 AI 能力的工程化实现，形成可测试闭环。",
            "前端开发页面与交互；后端开发接口、数据处理、权限校验和业务写入；Agent 开发工作流、Prompt、工具调用、结果解析、日志和灰度开关。",
            "可测试版本、接口实现、Agent 配置、Prompt 版本记录、调用日志。",
            "主流程可跑通；AI 能力可稳定调用；关键行为和异常可追踪。",
        ),
        (
            "6. 功能测试及数据验证",
            "同时验证普通功能质量、AI 输出质量、数据准确性和权限安全。",
            "开展功能测试、异常测试、权限测试、性能测试；使用样例集评估 AI 输出准确性、完整性、稳定性和业务可接受度；验证数据读取、写入、脱敏和埋点结果。",
            "测试报告、AI 评测表、数据验证记录、问题清单、上线建议。",
            "功能无阻塞问题；AI 输出达到上线标准；无明显越权、幻觉或高风险数据问题。",
        ),
        (
            "7. 上线、监控与复盘",
            "观察真实用户使用效果，将项目经验沉淀为下一次可复用的方法。",
            "灰度上线并观察使用率、采纳率、编辑率、失败率、成本和响应速度；收集用户反馈；复盘产品效果、技术问题、测试遗漏和协作问题。",
            "上线记录、效果看板、用户反馈、复盘报告、下一版优化项、SOP 更新项。",
            "上线效果可量化；问题可追踪；经验能沉淀到后续项目和标准文档中。",
        ),
    ]

    for title_text, goal, actions, outputs, gate in stages:
        if title_text.startswith("5. "):
            doc.add_page_break()
        add_heading(doc, title_text, 2)
        rows = [
            ("阶段目标", goal),
            ("关键动作", actions),
            ("主要产出", outputs),
            ("通过标准", gate),
        ]
        add_table(doc, ["项目", "内容"], rows, [3.0, 13.5], font_size=9.5)

    doc.add_page_break()
    add_heading(doc, "三、角色分工建议", 1)
    roles = [
        ("产品经理", "负责需求判断、竞品调研、产品方案、PRD、AI_Spec、验收标准和上线复盘。"),
        ("前端开发", "负责 AI 能力入口、交互状态、结果展示、用户确认、编辑、重试和异常反馈体验。"),
        ("后端开发", "负责接口、数据读取与写入、权限校验、业务规则、日志、灰度、开关和回滚能力。"),
        ("Agent/算法开发", "负责 Agent 工作流、Prompt、模型选择、工具调用、输出 schema、评测优化和版本管理。"),
        ("测试", "负责功能测试、权限测试、异常测试、AI 样例评测、数据验证和上线质量判断。"),
        ("运营/业务方", "负责真实场景反馈、试点用户组织、效果观察、案例沉淀和后续推广建议。"),
    ]
    add_table(doc, ["角色", "主要职责"], roles, [3.2, 13.3], font_size=9.5)

    add_heading(doc, "四、关键交付物清单", 1)
    deliverables = [
        ("需求分析文档", "阶段 1", "说明为什么做、为谁做、解决什么问题，以及是否适合用 AI。"),
        ("竞品调研报告", "阶段 2", "提供外部参考、体验基准和差异化依据。"),
        ("PRD / 原型", "阶段 3", "定义产品流程、页面交互、业务规则和用户操作闭环。"),
        ("AI_Spec", "阶段 3-4", "定义 AI 能力边界、输入输出、上下文、Agent 行为和评测标准。"),
        ("Agent 设计文档 / 技术方案", "阶段 4", "定义 Agent 工作流、模型、工具、接口、权限、日志、成本和回滚。"),
        ("测试与复盘报告", "阶段 6-7", "验证上线质量，并沉淀后续优化经验。"),
    ]
    add_table(doc, ["交付物", "所属阶段", "用途"], deliverables, [4.0, 2.6, 9.9], font_size=9.3)

    add_heading(doc, "五、AI_Spec 建议目录", 1)
    add_paragraph(doc, "AI_Spec 是 AI 产品项目中最关键的协作文档，建议作为 PRD 的配套文档存在。它的目标是让产品、开发、测试对 AI 能力的边界和验收口径保持一致。", after=8)
    spec_rows = [
        ("1", "能力概述", "能力名称、目标用户、使用场景、业务价值。"),
        ("2", "输入与上下文", "用户输入、系统可读取数据、数据范围、权限边界、上下文裁剪规则。"),
        ("3", "输出与交互", "输出格式、展示方式、用户确认/编辑/撤销/重试机制。"),
        ("4", "Agent 行为", "工作流、Prompt、模型选择、工具调用、输出 schema、失败兜底。"),
        ("5", "安全与权限", "哪些数据可读、哪些动作可执行、哪些结果必须人工确认。"),
        ("6", "评测与验收", "样例集、评分维度、通过标准、典型坏案例和上线门槛。"),
        ("7", "监控与迭代", "埋点、日志、采纳率、失败率、成本、反馈入口和版本记录。"),
    ]
    add_table(doc, ["序号", "模块", "说明"], spec_rows, [1.2, 3.6, 11.7], font_size=9.5)

    add_heading(doc, "六、贯穿原则", 1)
    principles = [
        ("先判断价值，再进入方案", "AI 项目不要因为“能做”就做，而要先证明场景真实、价值明确、失败可控。"),
        ("数据准备前置", "在产品方案阶段就明确数据来源、结构、权限和测试样例，避免开发后期才发现数据不可用。"),
        ("Agent 行为可解释、可追踪", "Prompt、模型、工具调用、输出 schema 和版本都应有记录，方便调试和复盘。"),
        ("测试覆盖 AI 输出质量", "AI 产品不仅要测功能，还要测准确性、完整性、稳定性、幻觉、越权和业务可接受度。"),
        ("上线后持续复盘", "通过真实数据观察使用率、采纳率、失败率和成本，把每个项目变成下一次项目的资产。"),
    ]
    add_table(doc, ["原则", "说明"], principles, [4.0, 12.5], font_size=9.5)

    add_heading(doc, "推荐项目表述", 3)
    add_paragraph(
        doc,
        "宏观 SOP 管项目怎么推进，微观 SOP 管一个 AI 能力怎么设计和验收。前者用于跨团队协作，后者沉淀在 PRD、AI_Spec、Agent 设计和测试评测中。",
        size=10.5,
        after=0,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
