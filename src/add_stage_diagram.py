from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SRC = PROJECT_ROOT / "build" / "AI产品项目标准推进SOP.docx"
OUT = PROJECT_ROOT / "docs" / "AI产品项目推进SOP.docx"
IMG = PROJECT_ROOT / "assets" / "ai_stage_flow.png"
OUT.parent.mkdir(parents=True, exist_ok=True)
IMG.parent.mkdir(parents=True, exist_ok=True)


FONT_CANDIDATES = [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simsun.ttc",
]


def font(size, bold=False):
    path = next((p for p in FONT_CANDIDATES if Path(p).exists()), None)
    if path:
        return ImageFont.truetype(path, size=size, index=0)
    return ImageFont.load_default()


def rounded_rect(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered(draw, box, text, fnt, fill, line_gap=8):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    lines = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            candidate = current + ch
            if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] - draw.textbbox((0, 0), line, font=fnt)[1] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        tw = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - tw) / 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def draw_arrow(draw, start, end, color="#6B7C93", width=5):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - direction * 18, y2 - 10), (x2 - direction * 18, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 10, y2 - direction * 18), (x2 + 10, y2 - direction * 18)]
    draw.polygon(pts, fill=color)


def make_diagram():
    W, H = 1800, 520
    img = Image.new("RGB", (W, H), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = font(32)
    sub_font = font(20)
    num_font = font(28)
    label_font = font(24)
    small_font = font(18)

    draw.text((70, 28), "AI 产品项目 7 阶段推进流程", font=title_font, fill="#0B2545")
    draw.text((72, 70), "前 4 阶段确认方向与可行性，后 3 阶段完成工程实现、验证和迭代", font=sub_font, fill="#5A6473")

    # Band backgrounds.
    rounded_rect(draw, (55, 112, 1745, 292), 22, "#F4F7FB", "#D7E2EE", 2)
    rounded_rect(draw, (55, 322, 1745, 485), 22, "#FFF9EC", "#E7D8A8", 2)
    draw.text((90, 130), "前 4 阶段：定义价值、设计方案、确认可行性", font=small_font, fill="#1F4D78")
    draw.text((90, 340), "后 3 阶段：工程实现、测试验证、上线迭代", font=small_font, fill="#7A5A00")

    top = [
        ("1", "需求分析\n与价值初判", "做不做"),
        ("2", "市场与\n竞品调研", "怎么做得更好"),
        ("3", "产品方案设计\n与数据准备", "怎么定义方案"),
        ("4", "Agent 设计\n与技术评审", "能不能落地"),
    ]
    bottom = [
        ("5", "前后端开发\n+ Agent 开发", "做出来"),
        ("6", "功能测试\n及数据验证", "验得过"),
        ("7", "上线、监控\n与复盘", "持续优化"),
    ]

    box_w, box_h = 320, 102
    top_y = 168
    top_xs = [100, 515, 930, 1345]
    bottom_y = 378
    bottom_xs = [1345, 930, 515]

    blue_fill = "#E8EEF5"
    blue_border = "#7FA6CC"
    gold_fill = "#FFF2CC"
    gold_border = "#D9B75C"

    top_boxes = []
    for (num, label, tag), x in zip(top, top_xs):
        box = (x, top_y, x + box_w, top_y + box_h)
        top_boxes.append(box)
        rounded_rect(draw, box, 20, blue_fill, blue_border, 3)
        draw.ellipse((x + 17, top_y + 20, x + 59, top_y + 62), fill="#2E74B5")
        draw_centered(draw, (x + 17, top_y + 20, x + 59, top_y + 62), num, num_font, "#FFFFFF", 0)
        draw_centered(draw, (x + 70, top_y + 14, x + box_w - 16, top_y + 72), label, label_font, "#0B2545", 5)
        draw_centered(draw, (x + 70, top_y + 72, x + box_w - 16, top_y + 94), tag, small_font, "#5A6473", 1)

    bottom_boxes = []
    for (num, label, tag), x in zip(bottom, bottom_xs):
        box = (x, bottom_y, x + box_w, bottom_y + box_h)
        bottom_boxes.append(box)
        rounded_rect(draw, box, 20, gold_fill, gold_border, 3)
        draw.ellipse((x + 17, bottom_y + 20, x + 59, bottom_y + 62), fill="#B98200")
        draw_centered(draw, (x + 17, bottom_y + 20, x + 59, bottom_y + 62), num, num_font, "#FFFFFF", 0)
        draw_centered(draw, (x + 70, bottom_y + 14, x + box_w - 16, bottom_y + 72), label, label_font, "#0B2545", 5)
        draw_centered(draw, (x + 70, bottom_y + 72, x + box_w - 16, bottom_y + 94), tag, small_font, "#5A6473", 1)

    # Top row arrows.
    for a, b in zip(top_boxes, top_boxes[1:]):
        draw_arrow(draw, (a[2] + 18, (a[1] + a[3]) // 2), (b[0] - 18, (b[1] + b[3]) // 2))
    # Snake transition: phase 4 goes down, then phases 5-7 flow right-to-left.
    draw_arrow(draw, ((top_boxes[-1][0] + top_boxes[-1][2]) // 2, top_boxes[-1][3] + 12), ((bottom_boxes[0][0] + bottom_boxes[0][2]) // 2, bottom_boxes[0][1] - 12))
    for a, b in zip(bottom_boxes, bottom_boxes[1:]):
        draw_arrow(draw, (a[0] - 18, (a[1] + a[3]) // 2), (b[2] + 18, (b[1] + b[3]) // 2))

    draw.text((75, 492), "阶段关口：需求评审 → 方案评审 → 技术评审 → 测试验收 → 上线复盘", font=small_font, fill="#5A6473")
    img.save(IMG, quality=95)


def paragraph_after(paragraph, text=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    paragraph._p.addnext(new_para._p)
    if text:
        new_para.add_run(text)
    return new_para


def set_alt_text(inline_shape, title, descr):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", descr)


def add_caption(paragraph, text):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 100, 115)


def insert_diagram():
    copyfile(SRC, OUT)
    doc = Document(str(OUT))
    target = None
    for p in doc.paragraphs:
        if "建议将 AI 产品项目统一拆成 7 个阶段" in p.text:
            target = p
            break
    if target is None:
        raise RuntimeError("Could not find insertion point")

    pic_para = paragraph_after(target)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.space_before = Pt(8)
    pic_para.paragraph_format.space_after = Pt(2)
    run = pic_para.add_run()
    shape = run.add_picture(str(IMG), width=Inches(6.45))
    set_alt_text(shape, "AI 产品项目 7 阶段推进流程图", "从需求分析、竞品调研、产品方案、Agent 设计、开发、测试到上线复盘的 7 阶段流程")

    caption_para = paragraph_after(pic_para)
    add_caption(caption_para, "图 1：AI 产品项目 7 阶段推进流程")

    break_para = paragraph_after(caption_para)
    break_para.add_run().add_break(WD_BREAK.PAGE)

    for p in doc.paragraphs:
        if p.text.strip() == "四、历史经验":
            p.paragraph_format.page_break_before = True
            break

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    make_diagram()
    insert_diagram()
