from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT = PROJECT_ROOT / "docs" / "source" / "AI产品项目复盘指南.docx"
DIAGRAM = PROJECT_ROOT / "assets" / "ai_lifecycle_flywheel.png"
OUT.parent.mkdir(parents=True, exist_ok=True)
DIAGRAM.parent.mkdir(parents=True, exist_ok=True)

FONT_EA = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 100, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F6F8FA"
PALE_BLUE = "F4F7FB"
PALE_GOLD = "FFF7E6"
BORDER = "D9E2EC"


def font_path():
    for path in [r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\simsun.ttc"]:
        if Path(path).exists():
            return path
    return None


def pil_font(size):
    path = font_path()
    if path:
        return ImageFont.truetype(path, size=size, index=0)
    return ImageFont.load_default()


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT_EA
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(p, before=0, after=6, line=1.12):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def para(doc, text="", size=11, color=None, bold=False, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para_spacing(p, before, after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {min(level, 3)}"
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    set_para_spacing(p, before, after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def cell_text(cell, text, size=9.3, bold=False, color=None, fill=None, align=None):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    set_para_spacing(p, after=0, line=1.14)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_widths(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = sum(int(Cm(w).twips) for w in widths_cm)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(width).twips)))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(Cm(width).twips)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths_cm, size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table_borders(table)
    set_table_widths(table, widths_cm)
    repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, h in enumerate(headers):
        cell_text(table.rows[0].cells[idx], h, size=9.4, bold=True, color=NAVY, fill=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        tr = table.add_row()
        set_row_cant_split(tr)
        for idx, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            fill = "FFFFFF" if len(table.rows) % 2 else "FBFCFE"
            cell_text(tr.cells[idx], val, size=size, fill=fill, align=align)
    para(doc, "", after=2)
    return table


def callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table_borders(table, color="CAD6E2", size="6")
    set_table_widths(table, [16.5])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    set_para_spacing(p2, after=0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    para(doc, "", after=2)


def add_rule(p, color="D7DBE2", size="8"):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EA
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        style = doc.styles[name]
        style.font.name = FONT_EA
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("AI 产品项目复盘")
    set_run_font(r, size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("AI 产品项目协作参考")
    set_run_font(r, size=9, color=GRAY)
    return doc


def draw_center(draw, box, text, fnt, fill, gap=4):
    x1, y1, x2, y2 = box
    max_w = x2 - x1 - 18
    lines = []
    for raw in text.split("\n"):
        cur = ""
        for ch in raw:
            cand = cur + ch
            if draw.textbbox((0, 0), cand, font=fnt)[2] <= max_w:
                cur = cand
            else:
                if cur:
                    lines.append(cur)
                cur = ch
        if cur:
            lines.append(cur)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] - draw.textbbox((0, 0), line, font=fnt)[1] for line in lines]
    total_h = sum(heights) + gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        tw = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - tw) / 2, y), line, font=fnt, fill=fill)
        y += h + gap


def arrow(draw, start, end, color="#6B7C93", width=5):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    if abs(x2 - x1) >= abs(y2 - y1):
        d = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - d * 17, y2 - 9), (x2 - d * 17, y2 + 9)]
    else:
        d = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - d * 17), (x2 + 9, y2 - d * 17)]
    draw.polygon(pts, fill=color)


def make_lifecycle_diagram():
    W, H = 1800, 560
    img = Image.new("RGB", (W, H), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title = pil_font(34)
    small = pil_font(19)
    label = pil_font(23)
    num = pil_font(24)

    draw.text((70, 30), "AI 产品全生命周期闭环", font=title, fill="#0B2545")
    draw.text((72, 74), "AI 产品不是一次性交付，而是上线后持续收集数据、优化 Prompt / Agent、重新评测并进入下一版本。", font=small, fill="#5A6473")

    stages = [
        ("1", "需求分析", "价值初判"),
        ("2", "竞品调研", "体验基准"),
        ("3", "产品方案", "数据准备"),
        ("4", "AI Spec", "边界与评测"),
        ("5", "Agent 设计", "Context / Tool"),
        ("6", "开发实现", "前后端 + Agent"),
        ("7", "测试验证", "Golden Standard"),
        ("8", "上线监控", "指标看板"),
        ("9", "Prompt 优化", "效果回流"),
        ("10", "下一版本", "持续迭代"),
    ]
    box_w, box_h = 260, 90
    xs_top = [70, 390, 710, 1030, 1350]
    xs_bottom = [1350, 1030, 710, 390, 70]
    y_top, y_bottom = 155, 355
    boxes = []
    for i, stage in enumerate(stages):
        x = xs_top[i] if i < 5 else xs_bottom[i - 5]
        y = y_top if i < 5 else y_bottom
        fill = "#E8EEF5" if i < 5 else "#FFF2CC"
        border = "#7FA6CC" if i < 5 else "#D9B75C"
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill=fill, outline=border, width=3)
        draw.ellipse((x + 16, y + 22, x + 54, y + 60), fill="#2E74B5" if i < 5 else "#B98200")
        draw_center(draw, (x + 16, y + 22, x + 54, y + 60), stage[0], num, "#FFFFFF")
        draw_center(draw, (x + 64, y + 14, x + box_w - 12, y + 55), stage[1], label, "#0B2545")
        draw_center(draw, (x + 64, y + 56, x + box_w - 12, y + 82), stage[2], small, "#5A6473")
        boxes.append((x, y, x + box_w, y + box_h))
    for a, b in zip(boxes[:4], boxes[1:5]):
        arrow(draw, (a[2] + 15, (a[1] + a[3]) // 2), (b[0] - 15, (b[1] + b[3]) // 2))
    arrow(draw, ((boxes[4][0] + boxes[4][2]) // 2, boxes[4][3] + 14), ((boxes[5][0] + boxes[5][2]) // 2, boxes[5][1] - 14))
    for a, b in zip(boxes[5:9], boxes[6:10]):
        arrow(draw, (a[0] - 15, (a[1] + a[3]) // 2), (b[2] + 15, (b[1] + b[3]) // 2))

    # Feedback loop from next version back to demand.
    draw.arc((68, 102, 1610, 535), start=22, end=338, fill="#9AA9BB", width=4)
    draw.polygon([(95, 155), (118, 145), (112, 170)], fill="#9AA9BB")
    draw.text((765, 506), "线上数据 -> 重新评测 -> 下一版本", font=small, fill="#5A6473")
    img.save(DIAGRAM, quality=95)


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=6, after=2)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.5))
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", "AI 产品从需求、方案、AI Spec、Agent、开发、测试、上线监控到下一版本的持续优化闭环")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cap, before=2, after=8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=GRAY)


def build():
    make_lifecycle_diagram()
    doc = setup_doc()

    para(doc, "项目复盘与协作规范", size=11, color=BLUE, bold=True, after=6)
    title_p = para(doc, "AI 产品项目复盘与全生命周期研发 SOP", size=23, color=NAVY, bold=True, after=4)
    sub = para(doc, "基于多个 AI 产品项目的实践经验沉淀", size=12.5, color=GRAY, after=12)
    add_rule(sub)

    add_table(
        doc,
        ["项目", "说明"],
        [
            ("文档版本", "v0.2"),
            ("适用场景", "AI 工作流、Agent 应用与文档自动化"),
            ("涉及角色", "产品经理、前端开发、后端开发、Agent/算法开发、测试、运营及项目协作方"),
            ("定位", "既是项目复盘，也是后续 AI 产品项目的协作 SOP 与方法论草案。"),
        ],
        [3.0, 13.5],
        size=9.8,
    )

    callout(
        doc,
        "核心共识",
        "AI 产品项目不是普通功能开发外接一个模型接口，而是同时管理产品体验、数据上下文、Agent 行为、工程稳定性和效果评测的协作过程。AI 产品本质是持续优化系统，而不是一次性交付。",
    )

    heading(doc, "一、为什么要沉淀这套 AI SOP", 1)
    para(doc, "在多个 AI 项目的实践中，项目类型虽然不同，但推进过程会暴露出一些共性问题。沉淀 SOP 的目的，不是增加流程负担，而是把已经踩过的坑变成团队后续协作的默认动作。", after=8)
    add_table(
        doc,
        ["共性问题", "对项目的影响", "SOP 中的对应机制"],
        [
            ("产品方案确定后才发现数据拿不到或数据质量不稳定", "方案返工，Agent 效果难以达到预期", "数据准备前置、AI Spec 明确上下文和数据边界"),
            ("Prompt 到开发阶段频繁调整", "产品、Agent、开发之间反复拉齐，成本高", "Prompt / Agent 版本管理，先用 Demo 和 Golden Standard 评估"),
            ("AI 输出没有统一评测标准", "上线判断依赖主观感受，难以复盘", "建立 Golden Standard、人工评测表和通过标准"),
            ("测试阶段只测功能，没有测 AI 效果", "功能可用但结果不可用，用户信任度不足", "功能测试 + AI 输出质量评测 + 数据验证并行"),
            ("上线以后缺少效果监控", "无法判断是否值得继续投入和优化", "上线指标看板、采纳率/失败率/成本监控、复盘机制"),
        ],
        [4.0, 5.6, 6.9],
        size=8.7,
    )

    heading(doc, "二、AI 产品全生命周期闭环", 1)
    para(doc, "建议将 AI 产品项目统一拆成 7 个主阶段，并在上线后进入数据回流和下一版本优化。前 4 个阶段解决“做什么、为什么做、怎么做”的问题，后 3 个阶段解决“如何实现、如何验证、如何持续优化”的问题。", after=6)
    add_picture(doc, DIAGRAM, "图 1：AI 产品全生命周期闭环")
    add_table(
        doc,
        ["阶段", "名称", "核心目标", "主要产出", "通过标准"],
        [
            ("1", "需求分析与价值初判", "判断需求是否真实、是否值得做、AI 是否能做、AI 是否适合做", "需求分析文档、价值判断、AI 适配度结论", "场景真实、价值明确、AI 有必要性"),
            ("2", "市场与竞品调研", "明确外部解法、体验基准和差异化方向", "竞品调研报告、可借鉴点、差异化机会", "知道别人怎么做，也知道我们应该怎样做"),
            ("3", "产品方案设计与数据准备", "把需求转化为可开发、可验证的产品方案", "原型、PRD、AI_Spec、开发/测试/验证数据、Golden Standard、数据埋点方案", "产品流程清楚，数据边界清楚，开发能理解并评估"),
            ("4", "Agent 设计与技术方案评审", "确认 AI 能力如何实现，以及工程上是否可行", "Agent Demo、Agent 设计文档、技术方案、接口文档、排期评估", "方案可实现，风险可控，成本可接受"),
            ("5", "产品前后端开发 + Agent 开发", "完成功能闭环和 AI 能力实现", "可测试版本、接口实现、Agent/Prompt 配置、调用日志、灰度/开关配置", "主流程跑通，AI 能力可调用，关键日志可追踪"),
            ("6", "功能测试及数据验证", "验证功能可用、AI 输出可信、数据安全", "测试报告、数据验证记录、AI 评测表、人工评测报告、问题清单、优化建议、上线建议", "功能稳定，AI 输出达标，无高风险问题"),
            ("7", "上线、监控与复盘", "观察真实使用效果，并沉淀经验", "上线记录、指标看板、用户反馈、复盘报告、Prompt/Agent 优化记录", "有数据可看，有问题可追，有经验可复用"),
        ],
        [1.0, 3.3, 4.0, 4.3, 3.9],
        size=7.8,
    )

    heading(doc, "三、AI Spec 标准模板", 1)
    para(doc, "AI_Spec 是 PRD 的配套文档，用于定义 AI 能力边界、输入输出、上下文、模型、工具、约束和评测方式。它让产品、Agent、开发、测试对同一件事形成可执行共识。", after=6)
    add_table(
        doc,
        ["模块", "需要说明的内容"],
        [
            ("目标", "AI 需要完成什么任务，服务什么用户场景，输出对业务有什么价值。"),
            ("输入 Input", "用户输入、业务上下文、历史数据、工具返回、上下文裁剪规则。"),
            ("输出 Output", "输出格式、JSON Schema、字段定义、异常输出、用户确认/编辑/撤销机制。"),
            ("模型 Model", "模型选择、调用方式、备用模型、模型升级策略、成本和延迟预估。"),
            ("Prompt", "System Prompt、User Prompt、Few-shot 示例、Prompt 版本记录。"),
            ("Tool", "Function Calling、MCP、知识库、业务接口、工具权限和失败处理。"),
            ("约束 Guardrail", "禁止输出、权限控制、Token 限制、敏感数据处理、人工确认边界。"),
            ("评测 Evaluation", "Golden Standard、准确性、完整性、稳定性、Latency、Cost、通过标准。"),
        ],
        [4.0, 12.5],
        size=9.1,
    )

    heading(doc, "四、Golden Standard 建设标准", 1)
    para(doc, "Golden Standard 是 AI 项目的核心资产，用于回答“AI 输出到底好不好”。它不是测试阶段临时造几条样例，而是在方案阶段就开始建设，并在每次 Prompt 修改、模型升级或 Agent 流程调整后重新评测。", after=6)
    add_table(
        doc,
        ["组成", "建议标准"],
        [
            ("真实 Case", "优先来自真实用户场景，早期可从 30-50 条起步，成熟能力逐步沉淀到 100-500 条。"),
            ("标准答案", "为每个 Case 标注理想输出、关键事实、禁止出现的错误和可接受变体。"),
            ("评分规则", "至少包含准确性、完整性、稳定性、业务可接受度、幻觉/越权风险。"),
            ("人工标注", "由产品、业务方、测试共同参与，避免只从技术视角判断。"),
            ("触发时机", "上线前、每次 Prompt 修改、每次模型升级、每次工具或上下文逻辑调整后都重新跑评测。"),
        ],
        [3.5, 13.0],
        size=9.1,
    )

    heading(doc, "五、角色分工建议", 1)
    add_table(
        doc,
        ["角色", "主要职责"],
        [
            ("项目经理", "负责项目总体进度、成本、质量管理，协调资源、拉齐团队认知。"),
            ("产品经理", "负责需求判断、竞品调研、产品方案、PRD、AI_Spec、Agent Demo 设计、验收标准、Agent 评测方案。"),
            ("前端开发", "负责 AI 能力入口、交互状态、结果展示、用户确认、编辑、重试和异常反馈体验。"),
            ("后端开发", "负责接口、数据读取与写入、权限校验、业务规则、日志、灰度、开关和回滚能力。"),
            ("Agent 开发", "负责 Agent 工作流、Context 管理、Memory、Long Context、Tool Routing、Fallback、Prompt、模型选择、输出 Schema、评测优化和版本管理。"),
            ("测试", "负责功能测试、权限测试、异常测试、AI 样例评测、数据验证和上线质量判断。"),
            ("运营/业务方", "负责真实场景反馈、试点用户组织、效果观察、案例沉淀和后续推广建议。"),
        ],
        [3.2, 13.3],
        size=9.1,
    )

    heading(doc, "六、关键交付物建议", 1)
    add_table(
        doc,
        ["交付物", "所属阶段", "用途"],
        [
            ("需求分析文档", "阶段 1", "说明为什么做、为谁做、解决什么问题，以及是否适合用 AI。"),
            ("竞品调研报告", "阶段 2", "提供外部参考、体验基准和差异化依据。"),
            ("PRD / 原型", "阶段 3", "定义产品流程、页面交互、业务规则和用户操作闭环。"),
            ("Golden Standard / 评测集", "阶段 3", "作为标准用于评估 AI 输出质量，支撑上线前和迭代后的效果判断。"),
            ("AI_Spec", "阶段 3-4", "定义 AI 能力边界、输入输出、上下文、Agent 行为和评测标准。"),
            ("Agent 设计文档 / 技术方案", "阶段 4", "定义 Agent 工作流、模型、工具、接口、权限、日志、成本和回滚。"),
            ("埋点与监控方案", "阶段 3/7", "定义上线后要看的调用率、采纳率、失败率、成本、延迟等指标。"),
            ("测试与复盘报告", "阶段 6-7", "验证上线质量，并沉淀后续优化经验。"),
        ],
        [4.0, 2.4, 10.1],
        size=8.9,
    )

    heading(doc, "七、AI 项目风险治理", 1)
    add_table(
        doc,
        ["风险", "典型表现", "应对机制"],
        [
            ("Prompt 失效", "同类输入输出不稳定，修改一处影响其他场景", "Prompt 版本管理、Golden Standard 回归评测"),
            ("模型升级效果下降", "模型替换后准确率、风格或工具调用行为变化", "A/B 评测、灰度切换、回滚策略"),
            ("Token 成本上涨", "调用量上升或上下文过长导致成本不可控", "成本监控、上下文裁剪、缓存和分层模型策略"),
            ("幻觉", "输出不存在的事实、任务或结论", "引用依据、结构化输出、Golden Standard、人工确认"),
            ("延迟过高", "用户等待时间长，影响转化和采纳", "Streaming、异步任务、缓存、模型分层"),
            ("工具调用失败", "Agent 调接口失败、参数错误或重复执行", "Retry、幂等设计、Fallback、错误日志和告警"),
            ("越权或数据泄露", "读取了用户无权访问的数据或暴露敏感信息", "服务端权限校验、数据脱敏、审计日志"),
        ],
        [3.0, 5.5, 8.0],
        size=8.8,
    )

    heading(doc, "八、上线监控指标", 1)
    para(doc, "AI 产品上线后要持续看“有没有人用、结果能不能用、成本是否可控、哪里需要优化”。建议在灰度阶段就建立效果看板。", after=6)
    add_table(
        doc,
        ["指标", "说明"],
        [
            ("AI 调用率", "目标用户中有多少人实际触发 AI 能力。"),
            ("AI 采纳率", "AI 输出被直接采纳或部分采纳的比例。"),
            ("一次成功率", "用户一次调用后即可完成目标任务的比例。"),
            ("Retry 率", "用户重复生成或重新调用的比例，用于判断输出稳定性。"),
            ("用户编辑率", "用户对 AI 输出进行修改的比例，用于判断可用性和可信度。"),
            ("失败率", "模型调用、工具调用、解析、权限等失败比例。"),
            ("人工介入率", "需要人工确认、修正或兜底的比例。"),
            ("Latency", "端到端响应耗时，尤其关注 P50 / P90 / P95。"),
            ("Token 成本", "单次调用成本、总成本、按用户/功能/模型拆分的成本。"),
        ],
        [4.0, 12.5],
        size=9.0,
    )

    doc.add_page_break()
    heading(doc, "九、关键原则建议", 1)
    add_table(
        doc,
        ["原则", "说明"],
        [
            ("先判断价值，再进入方案", "AI 项目不要因为“能做”就做，而要先证明场景真实、价值明确、失败可控。"),
            ("数据准备前置", "在产品方案阶段就明确数据来源、结构、权限和测试样例，避免开发后期才发现数据不可用。"),
            ("AI Spec 先于开发", "在进入开发前先定义输入、输出、上下文、工具、约束和评测标准。"),
            ("Context 比 Prompt 更关键", "复杂 AI 项目中，效果往往取决于上下文管理、工具路由和数据质量，而不只是 Prompt 文案。"),
            ("测试覆盖 AI 输出质量", "AI 产品不仅要测功能，还要测准确性、完整性、稳定性、幻觉、越权和业务可接受度。"),
            ("上线后持续复盘", "通过真实数据观察使用率、采纳率、失败率和成本，把每个项目变成下一次项目的资产。"),
        ],
        [4.2, 12.3],
        size=9.0,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
